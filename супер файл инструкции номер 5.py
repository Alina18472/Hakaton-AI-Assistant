import os
import pandas as pd
import torch
import numpy as np
from transformers import DistilBertTokenizer, DistilBertForSequenceClassification, Trainer, TrainingArguments
from sentence_transformers import SentenceTransformer
from torch.utils.data import Dataset
from sklearn.model_selection import train_test_split
from sklearn.utils.class_weight import compute_class_weight
from torch.nn import CrossEntropyLoss
from docx import Document
import subprocess  # For opening .docx files

# Load dataset with requests and labels
data = pd.read_excel('dataset.xlsx')
requests = data['Topic'].tolist()
services = data['label'].astype('category').cat.codes.tolist()
request_ids = data['№'].tolist()  # Assuming 'RequestID' column has unique request numbers
labels_to_services = data['label'].astype('category').cat.categories

# Load Tokenizer and Embedding Model
tokenizer = DistilBertTokenizer.from_pretrained('distilbert-base-uncased')
sbert_model = SentenceTransformer('all-MiniLM-L6-v2')

# Load and process detailed Word files
detailed_texts = []
detailed_files_path = 'Instructions'  # Path to the folder with Word files

for filename in os.listdir(detailed_files_path):
    if filename.endswith(".docx"):
        doc = Document(os.path.join(detailed_files_path, filename))
        detailed_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])  # Combine all paragraphs
        detailed_texts.append((filename, detailed_text))  # Store filename and content

# Convert detailed texts into embeddings for comparison
detailed_embeddings = sbert_model.encode([text for _, text in detailed_texts])

# Convert problem requests into embeddings for comparison
request_embeddings = sbert_model.encode(requests)

# Find best matches for each problem request
request_to_instructions = {}
top_n = 3  # Number of top matches to find

for i, request_embedding in enumerate(request_embeddings):
    similarities = np.dot(detailed_embeddings, request_embedding)
    best_match_indices = np.argsort(similarities)[-top_n-1:][::-1]  # Get top N matches (plus the best match)
    request_to_instructions[request_ids[i]] = [(detailed_texts[idx][0], detailed_texts[idx][1]) for idx in best_match_indices]

# Function to create and open a docx file
def create_docx(filename, matches):
    doc = Document()
    doc.add_heading('Matching Instructions', level=1)
    for match_filename, match_text in matches:
        doc.add_heading(match_filename, level=2)
        doc.add_paragraph(match_text)
    doc.save(filename)
    subprocess.Popen(["start", filename], shell=True)  # For Windows, use "open" for macOS

# Example: Create and open a document for a specific request ID
example_request_id = request_ids[0]  # Change as needed
if example_request_id in request_to_instructions:
    matches = request_to_instructions[example_request_id]
    create_docx(f'{example_request_id}_matches.docx', matches)
