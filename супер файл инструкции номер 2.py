import torch
from transformers import DistilBertTokenizer, DistilBertForSequenceClassification
from sentence_transformers import SentenceTransformer, util
import pandas as pd
import tkinter as tk
from tkinter import Toplevel, Text
import os
import docx  # Library to handle Word documents

# Load data and create service names
data = pd.read_excel('dataset.xlsx')
texts = data['Topic'].tolist()
sbert_model = SentenceTransformer('all-MiniLM-L6-v2')
text_embeddings = sbert_model.encode(texts, convert_to_tensor=True)
service_names = list(data['label'].astype('category').cat.categories)

# Load the trained model and tokenizer
tokenizer = DistilBertTokenizer.from_pretrained("enhanced_trained_model")
model = DistilBertForSequenceClassification.from_pretrained("enhanced_trained_model")

# Instructions from dataset
instructions = data['Solution'].dropna().tolist()
instructions = [str(instruction) for instruction in instructions]
instruction_embeddings = sbert_model.encode(instructions, convert_to_tensor=True)

# Function to find and display matching Word content
def show_detailed_instruction(instruction_text):
    folder_path = 'Instructions'  # Path where Word files are stored
    best_file = None
    best_score = -1

    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            doc_path = os.path.join(folder_path, filename)
            doc = docx.Document(doc_path)
            full_text = ' '.join([p.text for p in doc.paragraphs])

            # Calculate similarity with the instruction text
            doc_embedding = sbert_model.encode(full_text, convert_to_tensor=True)
            instruction_embedding = sbert_model.encode(instruction_text, convert_to_tensor=True)
            score = util.pytorch_cos_sim(instruction_embedding, doc_embedding).item()
            

            if score > best_score:
                best_score = score
                best_file = full_text

    if best_file:
        # Open a new window to display the content
        top = Toplevel(window)
        top.title("Detailed Instruction")
        text_widget = Text(top, wrap='word')
        text_widget.insert('1.0', best_file)
        text_widget.pack(expand=True, fill='both')

# Function to analyze the problem and display results
import torch
from transformers import DistilBertTokenizer, DistilBertForSequenceClassification
from sentence_transformers import SentenceTransformer, util
import pandas as pd
import tkinter as tk
from tkinter import Toplevel, Text
import os
import docx  # Library to handle Word documents

# Load data and create service names
data = pd.read_excel('dataset.xlsx')
texts = data['Topic'].tolist()
sbert_model = SentenceTransformer('all-MiniLM-L6-v2')
text_embeddings = sbert_model.encode(texts, convert_to_tensor=True)
service_names = list(data['label'].astype('category').cat.categories)

# Load the trained model and tokenizer
tokenizer = DistilBertTokenizer.from_pretrained("enhanced_trained_model")
model = DistilBertForSequenceClassification.from_pretrained("enhanced_trained_model")

# Instructions from dataset
instructions = data['Solution'].dropna().tolist()
instructions = [str(instruction) for instruction in instructions]
instruction_embeddings = sbert_model.encode(instructions, convert_to_tensor=True)

# Function to find and display the most similar Word file content
def show_detailed_instruction(instruction_text):
    folder_path = 'Instructions'  # Path where Word files are stored
    best_file_text = None
    best_score = -1

    # Encode the brief instruction text to find the most similar Word file
    instruction_embedding = sbert_model.encode(instruction_text, convert_to_tensor=True)

    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            doc_path = os.path.join(folder_path, filename)
            doc = docx.Document(doc_path)
            full_text = []

            # Collect paragraphs and check for multiple empty lines
            for para in doc.paragraphs:
                if para.text.strip():  # Только непустые строки
                    full_text.append(para.text.strip())
                elif full_text and full_text[-1]:  # Одна пустая строка разрешена
                    full_text.append("") 

            # Join paragraphs with single line breaks between sections
            full_text = '\n\n'.join(full_text)

            # Calculate similarity with the brief instruction
            doc_embedding = sbert_model.encode(full_text, convert_to_tensor=True)
            score = util.pytorch_cos_sim(instruction_embedding, doc_embedding).item()
            
            if score > best_score:
                best_score = score
                best_file_text = full_text

    if best_file_text:
        # Create a new top-level window and a canvas for centering
        top = Toplevel(window)
        top.title("Detailed Instruction")
        
        # Create a canvas to center the Text widget
        canvas = tk.Canvas(top)
        canvas.pack(side="left", fill="both", expand=True)
        
        # Add a scrollbar to the canvas
        scrollbar = tk.Scrollbar(top, command=canvas.yview)
        scrollbar.pack(side="right", fill="y")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Create a frame on the canvas to hold the text widget
        frame = tk.Frame(canvas)
        canvas.create_window((0, 0), window=frame, anchor="center")

        # Create and configure the Text widget
        text_widget = Text(frame, wrap='word', font=("Arial", 12), spacing1=10, spacing2=5, spacing3=10)
        text_widget.insert('1.0', best_file_text)

        # Disable editing for the Text widget
        text_widget.config(state='disabled')
        text_widget.pack(expand=True, fill='both')

        # Update the scroll region and configure the canvas for centering effect
        frame.update_idletasks()
        canvas.configure(scrollregion=canvas.bbox("all"))



# Function to analyze the problem and display results
def analyze_problem(problem_text):
    # Tokenization and classification
    inputs = tokenizer(problem_text, return_tensors="pt", truncation=True, padding="max_length", max_length=64)
    outputs = model(**inputs)
    softmaxed = torch.nn.functional.softmax(outputs.logits, dim=-1)
    top_probs, top_indices = torch.topk(softmaxed, k=3)  # Top 3 predictions

    # Extract top services and probabilities
    try:
        top_services = [(service_names[idx], prob.item() * 100) for idx, prob in zip(top_indices[0], top_probs[0])]
    except IndexError:
        clear_results()
        tk.Label(result_frame, text="Error: Unable to retrieve service names from predictions.").pack()
        return

    # Find multiple similar problems
    query_embedding = sbert_model.encode(problem_text, convert_to_tensor=True)
    cos_sim = util.pytorch_cos_sim(query_embedding, text_embeddings)
    top_similarities, top_indices = torch.topk(cos_sim[0], k=min(5, len(texts)))  # Get top 5 similar problems

    # Extract the most similar problems and their similarity percentages
    top_problems = []
    for idx, sim in zip(top_indices, top_similarities):
        similar_problem = texts[idx.item()]
        similarity_percent = round(sim.item() * 100, 2)
        top_problems.append((similar_problem, similarity_percent))

    # Find the most similar brief instruction
    instruction_cos_sim = util.pytorch_cos_sim(query_embedding, instruction_embeddings)
    best_instruction_idx = torch.argmax(instruction_cos_sim).item()
    instruction = instructions[best_instruction_idx] if pd.notnull(instructions[best_instruction_idx]) else "No instruction found."

    # Clear previous results
    clear_results()

    # Display the results
    tk.Label(result_frame, text="Service suggestions:", font=("Arial", 14)).pack()
    for i, (service, prob) in enumerate(top_services[:3]):  # Show only the top 3 services
        tk.Label(result_frame, text=f"{i+1}. {service}: {prob:.2f}%", font=("Arial", 12)).pack()

    tk.Label(result_frame, text="\nMost similar problems:", font=("Arial", 14)).pack()
    for i, (problem, percentage) in enumerate(top_problems[:3]):  # Show only the top 3 similar problems
        tk.Label(result_frame, text=f"{i+1}. {problem} ({percentage}%)", font=("Arial", 12)).pack()

    tk.Label(result_frame, text="\nSuggested instruction:", font=("Arial", 14)).pack()
    tk.Label(result_frame, text=f"{instruction}", font=("Arial", 12)).pack()

    # Button to view detailed instruction
    detailed_button = tk.Button(result_frame, text="View Detailed Instruction", command=lambda: show_detailed_instruction(instruction))
    detailed_button.pack()

def clear_results():
    for widget in result_frame.winfo_children():
        widget.destroy()

# Tkinter GUI setup
window = tk.Tk()
window.title("Problem Analyzer")

label_problem = tk.Label(window, text="Enter your problem description:")
label_problem.pack()
entry_problem = tk.Entry(window, width=50)
entry_problem.pack()

button_analyze = tk.Button(window, text="Analyze", command=lambda: analyze_problem(entry_problem.get()))
button_analyze.pack()

label_result = tk.Label(window, text="Result:")
label_result.pack()

result_frame = tk.Frame(window)
result_frame.pack()

window.mainloop()


def clear_results():
    for widget in result_frame.winfo_children():
        widget.destroy()

# Tkinter GUI setup
window = tk.Tk()
window.title("Problem Analyzer")

label_problem = tk.Label(window, text="Enter your problem description:")
label_problem.pack()
entry_problem = tk.Entry(window, width=50)
entry_problem.pack()

button_analyze = tk.Button(window, text="Analyze", command=lambda: analyze_problem(entry_problem.get()))
button_analyze.pack()

label_result = tk.Label(window, text="Result:")
label_result.pack()

result_frame = tk.Frame(window)
result_frame.pack()

window.mainloop()
